import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkinter import font as tkFont
import os
import re
import sys

# --- INICIO: Añadir raíz del proyecto a sys.path y definir PROJECT_ROOT_DIR ---
script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(script_dir)
if project_root not in sys.path:
    sys.path.insert(0, project_root)
PROJECT_ROOT_DIR = project_root
# --- FIN ---

try:
    from PIL import Image, ImageEnhance
except ImportError: messagebox.showerror("Error Importación", "Pillow no instalado."); exit()
try:
    import pytesseract
except ImportError: messagebox.showerror("Error Importación", "pytesseract no instalado."); exit()
try:
    from docx import Document
except ImportError: messagebox.showerror("Error Importación", "python-docx no instalado."); exit()
try:
    from modules.tesseract_config import configure_tesseract
    print("INFO: Configurando Tesseract OCR...")
    tesseract_configurado_ok = configure_tesseract()
    if not tesseract_configurado_ok: print("ADVERTENCIA GUI: Config Tesseract OCR falló.")
    else: print("INFO GUI: Tesseract OCR configurado.")
except ImportError:
    print("ERROR CRÍTICO: No se importó 'configure_tesseract'.")
    messagebox.showerror("Error Módulo", "No se cargó config Tesseract. OCR no disponible.")
    tesseract_configurado_ok = False
except Exception as e_config_module:
    print(f"ERROR INESPERADO config Tesseract: {e_config_module}")
    messagebox.showerror("Error Configuración", f"Error config Tesseract:\n{e_config_module}")
    tesseract_configurado_ok = False

secciones_info = []
FONT_FAMILY = "Inter"
try:
    default_font_obj = tkFont.nametofont("TkDefaultFont"); DEFAULT_FONT_SIZE = default_font_obj.actual()["size"]
except: DEFAULT_FONT_SIZE = 10
FONT_NORMAL = (FONT_FAMILY, DEFAULT_FONT_SIZE); FONT_BOLD = (FONT_FAMILY, DEFAULT_FONT_SIZE, "bold")
FONT_PLACEHOLDER = (FONT_FAMILY, DEFAULT_FONT_SIZE); FONT_TOGGLE_ARROW = (FONT_FAMILY, DEFAULT_FONT_SIZE + 1)
FONT_TOGGLE_PASO_BOLD = (FONT_FAMILY, DEFAULT_FONT_SIZE + 2, "bold"); FONT_TOGGLE_TITLE_NORMAL = (FONT_FAMILY, DEFAULT_FONT_SIZE + 2)
FONT_MAIN_TITLE = (FONT_FAMILY, DEFAULT_FONT_SIZE + 6, "bold"); FONT_SECTION_HEADER = (FONT_FAMILY, DEFAULT_FONT_SIZE + 2, "bold")
FONT_SUB_HEADER = (FONT_FAMILY, DEFAULT_FONT_SIZE + 1, "bold"); FONT_BUTTON_ACTION_MAIN = (FONT_FAMILY, DEFAULT_FONT_SIZE, "bold") # Título del Word en Paso 2
COLOR_BOTON_ACCION_PRINCIPAL_AZUL = "#0078D4"; COLOR_TEXTO_BOTON_AZUL = "white"
COLOR_BOTON_DESHABILITADO_BG = "#cccccc"; COLOR_BOTON_DESHABILITADO_FG = "#666666"
COLOR_TEXTO_EXITO = "green"; COLOR_TEXTO_ADVERTENCIA = "#cc5500"
COLOR_GRIS_TEXTO_PLACEHOLDER = "#888888"; COLOR_GRIS_TEXTO_SECUNDARIO = "#595959"
COLOR_BOTON_MINIMALISTA_BG = "#f0f0f0"; COLOR_BOTON_MINIMALISTA_FG = "#333333"
COLOR_BOTON_MINIMALISTA_ACTIVE_BG = "#e0e0e0"; COLOR_BOTON_MINIMALISTA_BORDER = "#ababab"
COLOR_HEADER_LOCKED_BG = "#e8e8e8"; COLOR_HEADER_LOCKED_FG = "#a0a0a0"
COLOR_HEADER_ACTIVE_BG = "#e0e0e0"; COLOR_HEADER_ACTIVE_FG = "#333333"
PLACEHOLDER_INGRESE_VALOR = "Ingrese valor"; PLACEHOLDER_EJ_20 = "Ej: 20"; PLACEHOLDER_UNIDAD = "Unidad"

def preprocesar_imagen(ruta_imagen): # Sin cambios
    try:
        imagen = Image.open(ruta_imagen); imagen = imagen.convert('L')
        imagen = imagen.resize((imagen.width * 2, imagen.height * 2))
        enhancer = ImageEnhance.Contrast(imagen); imagen = enhancer.enhance(2)
        imagen = imagen.point(lambda x: 255 if x > 138 else 0, '1'); return imagen
    except FileNotFoundError: print(f"Error: No se pudo encontrar: '{ruta_imagen}'"); return None
    except Image.UnidentifiedImageError: print(f"Error: Archivo en '{ruta_imagen}' no es imagen válida."); return None
    except Exception as e: print(f"Error al preprocesar '{ruta_imagen}': {e}"); return None

def guardar_en_word(texto, nombre_archivo_word): # Usa PROJECT_ROOT_DIR
    try:
        docs_dir_abs = os.path.join(PROJECT_ROOT_DIR, "docs")
        if not os.path.exists(docs_dir_abs): os.makedirs(docs_dir_abs, exist_ok=True); print(f"Dir '{docs_dir_abs}' creado.")
        ruta_word = os.path.join(docs_dir_abs, nombre_archivo_word)
        documento = Document(); match_enunciado = re.search(r"#(\d+)", nombre_archivo_word)
        titulo_enunciado = f"Enunciado {match_enunciado.group(1)}" if match_enunciado else "Enunciado Procesado"
        documento.add_heading(titulo_enunciado, level=1) # Esto crea un estilo 'Heading 1' por defecto
        documento.add_paragraph(texto if texto else "No se pudo transcribir texto de la imagen.")
        documento.save(ruta_word); print(f"Texto guardado en: {ruta_word}"); return True
    except Exception as e: print(f"Error al guardar Word '{nombre_archivo_word}': {e}"); return False

def ejecutar_flujo_completo_paso1(ruta_imagen_seleccionada, section_data_paso1): # Modificado para guardar texto OCR
    if not ruta_imagen_seleccionada: messagebox.showerror("Error P1", "No imagen seleccionada."); return False
    print(f"P1: Preprocesando '{ruta_imagen_seleccionada}'..."); img_proc = preprocesar_imagen(ruta_imagen_seleccionada)
    if img_proc is None: messagebox.showerror("Error P1", "Fallo preprocesamiento."); return False
    print("P1: Imagen preprocesada."); print("P1: Realizando OCR...")
    texto_transcrito = "" # Inicializar
    try:
        texto_transcrito = pytesseract.image_to_string(img_proc, lang='spa')
        texto_transcrito = re.sub(r'-\n', '', texto_transcrito)
        print(f"P1: OCR OK. Texto(100): '{texto_transcrito[:100]}...'")
        if not texto_transcrito.strip(): messagebox.showwarning("OCR", "OCR no extrajo texto. Doc en blanco.")
        section_data_paso1['texto_ocr_obtenido'] = texto_transcrito # Guardar para Paso 2
    except pytesseract.TesseractNotFoundError: messagebox.showerror("Error Tesseract", "Tesseract no instalado/PATH."); return False
    except pytesseract.TesseractError as e:
        msg = f"Error Tesseract OCR: {e}"
        if "language 'spa' is not supported" in str(e).lower() or "error opening data file" in str(e).lower():
            msg = "Error: Paquete idioma 'spa' Tesseract no instalado."
        messagebox.showerror("Error Tesseract", msg); print(msg); return False
    except Exception as e: msg = f"Error OCR: {e}"; messagebox.showerror("Error OCR", msg); print(msg); return False
    nombre_doc = "WORD #1.docx"; print(f"P1: Guardando en '{nombre_doc}'...")
    if guardar_en_word(texto_transcrito, nombre_doc):
        print(f"P1: '{nombre_doc}' guardado."); return True
    else: messagebox.showerror("Error P1", f"No se pudo guardar '{nombre_doc}'."); return False

def leer_texto_desde_word(ruta_word): # Modificado para devolver (título, cuerpo)
    try:
        if not os.path.exists(ruta_word):
            return None, f"Archivo no encontrado: {os.path.basename(ruta_word)}"
        doc = Document(ruta_word)
        titulo_word = None
        parrafos_texto = []
        for para in doc.paragraphs:
            # Asumimos que el primer 'Heading 1' es el título principal.
            # 'python-docx' asigna estilos por defecto como 'Heading 1', 'Heading 2', etc.
            if titulo_word is None and para.style.name == 'Heading 1' and para.text.strip():
                titulo_word = para.text.strip()
            elif para.text.strip(): # Añadir al cuerpo si no es el título ya capturado
                if not (titulo_word and titulo_word == para.text.strip()): # Evitar duplicar el título en el cuerpo
                    parrafos_texto.append(para.text.strip())
        
        # Si no se encontró un 'Heading 1' pero hay párrafos, tomar el primero como título
        if titulo_word is None and parrafos_texto:
            titulo_word = parrafos_texto.pop(0) # Tomar y remover de párrafos
        elif titulo_word is None and not parrafos_texto: # Si no hay título ni párrafos
            return "Documento Vacío", "El archivo Word parece no contener texto."

        cuerpo_texto = "\n\n".join(parrafos_texto) # Usar doble salto de línea entre párrafos
        return titulo_word, cuerpo_texto
    except Exception as e:
        print(f"Error al leer Word '{ruta_word}': {e}")
        return None, f"Error al leer: {os.path.basename(ruta_word)} ({e})"

def cargar_y_mostrar_word_paso2(section_data):
    label_titulo_word = section_data.get('word_title_label')
    text_widget_word = section_data.get('word_text_widget') # Ahora es un tk.Text
    if not label_titulo_word or not text_widget_word: return

    nombre_archivo_word = "WORD #1.docx"
    ruta_completa_word = os.path.join(PROJECT_ROOT_DIR, "docs", nombre_archivo_word)
    print(f"P2: Leyendo '{ruta_completa_word}'...")
    titulo_leido, cuerpo_leido = leer_texto_desde_word(ruta_completa_word)

    if titulo_leido is not None:
        label_titulo_word.config(text=titulo_leido)
    else:
        label_titulo_word.config(text="No se pudo cargar el título del documento.")

    text_widget_word.config(state='normal') # Habilitar para modificar
    text_widget_word.delete('1.0', tk.END) # Limpiar contenido anterior
    if cuerpo_leido is not None:
        text_widget_word.insert(tk.END, cuerpo_leido)
        print(f"P2: Contenido de '{nombre_archivo_word}' cargado.")
    else: # cuerpo_leido es el mensaje de error o "documento vacío"
        text_widget_word.insert(tk.END, "No se pudo cargar el contenido del documento o está vacío.")
        print(f"P2: Fallo al cargar cuerpo de '{nombre_archivo_word}'.")
    text_widget_word.config(state='disabled') # Volver a solo lectura

def setup_placeholder(entry, placeholder_text): # Sin cambios
    style = ttk.Style(); original_fg='black'
    try:
        fg_lookup = style.lookup('TEntry', 'foreground');
        if fg_lookup and str(fg_lookup) != COLOR_GRIS_TEXTO_PLACEHOLDER : original_fg = fg_lookup
        else: fg_lookup_general = style.lookup('.', 'foreground');
        if fg_lookup_general and str(fg_lookup_general) != COLOR_GRIS_TEXTO_PLACEHOLDER : original_fg = fg_lookup_general
    except tk.TclError: pass
    except Exception: pass
    def on_focusin(event, c_entry=entry, txt=placeholder_text, o_fg=original_fg):
        if c_entry.get() == txt and str(c_entry.cget('foreground')) == COLOR_GRIS_TEXTO_PLACEHOLDER:
            c_entry.delete(0, tk.END); c_entry.config(foreground=o_fg)
    def on_focusout(event, c_entry=entry, txt=placeholder_text, o_fg=original_fg):
        if not c_entry.get().strip():
            c_entry.delete(0, tk.END); c_entry.insert(0, txt); c_entry.config(foreground=COLOR_GRIS_TEXTO_PLACEHOLDER)
    entry.insert(0, placeholder_text); entry.config(foreground=COLOR_GRIS_TEXTO_PLACEHOLDER)
    entry.bind("<FocusIn>", on_focusin); entry.bind("<FocusOut>", on_focusout)

def _actualizar_estilo_header(section_data): # Sin cambios
    header_frame = section_data.get('header_frame'); widgets_in_header = section_data.get('widgets_in_header', [])
    if not header_frame: return
    if section_data.get('is_unlocked', False):
        header_frame.config(style="HeaderFrame.TFrame", cursor="hand2")
        for widget in widgets_in_header:
            if widget != header_frame: widget.config(style="HeaderLabel.TLabel", cursor="hand2")
    else:
        header_frame.config(style="LockedHeaderFrame.TFrame", cursor="arrow")
        for widget in widgets_in_header:
            if widget != header_frame: widget.config(style="LockedHeaderLabel.TLabel", cursor="arrow")

def _actualizar_estado_visual_seccion(section_data): # Sin cambios
    arrow_widget = section_data.get('arrow_widget')
    if section_data['is_expanded'] and section_data.get('is_unlocked', False):
        section_data['frame_contenido'].pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5)) # fill=tk.BOTH, expand=True
        if arrow_widget: arrow_widget.config(text="▲")
    else:
        section_data['is_expanded'] = False; section_data['frame_contenido'].pack_forget()
        if arrow_widget: arrow_widget.config(text="▼")

def header_click_handler(event, section_data_clicked): # Modificado
    if section_data_clicked.get('is_unlocked', False):
        is_being_expanded = not section_data_clicked['is_expanded']
        section_data_clicked['is_expanded'] = is_being_expanded
        if is_being_expanded and section_data_clicked.get('paso_numero') == 2:
            cargar_y_mostrar_word_paso2(section_data_clicked)
        _actualizar_estado_visual_seccion(section_data_clicked)
    else: messagebox.showinfo("Paso Bloqueado", "Complete el paso anterior primero.")

def accion_principal_paso(section_data): # Modificado
    action_btn = section_data.get('action_button_principal');
    if not action_btn: return
    btn_text = action_btn.cget('text'); paso_titulo_completo = section_data.get('titulo_texto', "este paso")
    operacion_paso_especifico_exitosa = True
    mensaje_popup_exito = f"Acción '{btn_text}' completada para {paso_titulo_completo}."

    if btn_text == "PROCESAR IMAGEN":
        if not tesseract_configurado_ok: messagebox.showerror("Error Tesseract", "Tesseract no configurado."); return
        ruta_img = section_data.get('ruta_imagen_seleccionada')
        if not ruta_img: messagebox.showerror("Error P1", "No imagen seleccionada."); return
        if ejecutar_flujo_completo_paso1(ruta_img, section_data): # Pasar section_data
            mensaje_popup_exito = "WORD #1 creado con éxito."
        else: operacion_paso_especifico_exitosa = False; return
    elif btn_text == "INGRESAR DATOS AL MODELO":
        # ... (lógica de print sin cambios) ...
        print("\n--- Valores Ingresados en Paso 3 ---") # (código de print de Paso 3 va aquí)
        widgets_p3=section_data.get('widgets_contenido_a_deshabilitar',[]);
        try:
            if len(widgets_p3)>=6: l_v,l_u,m_v,m_u,k_v,s_v=widgets_p3[0].get(),widgets_p3[1].get(),widgets_p3[2].get(),widgets_p3[3].get(),widgets_p3[4].get(),widgets_p3[5].get();print(f" L:'{l_v if l_v!=PLACEHOLDER_EJ_20 else'(no)'}' {l_u}|Mu:'{m_v if m_v!=PLACEHOLDER_INGRESE_VALOR else'(no)'}' {m_u if m_u!=PLACEHOLDER_UNIDAD else'(no)'}|K:'{k_v if k_v!=PLACEHOLDER_INGRESE_VALOR else'(no)'}'|S:'{s_v if s_v!=PLACEHOLDER_INGRESE_VALOR else'(no)'}'")
            else: print(" Adv:No se leyeron campos P3.")
        except Exception as e: print(f" Err print P3:{e}")
        print("---------------------------------\n"); operacion_paso_especifico_exitosa = True
    
    if not operacion_paso_especifico_exitosa: return
    messagebox.showinfo("Proceso Completado", mensaje_popup_exito)
    action_btn.config(state="disabled", bg=COLOR_BOTON_DESHABILITADO_BG, fg=COLOR_BOTON_DESHABILITADO_FG)
    for btn_s in section_data.get('botones_secundarios_a_deshabilitar', []):
        btn_s.config(state="disabled", bg=COLOR_BOTON_DESHABILITADO_BG, fg=COLOR_BOTON_DESHABILITADO_FG)
    for widget in section_data.get('widgets_contenido_a_deshabilitar', []):
        try:
            if isinstance(widget,(ttk.Entry, ttk.Combobox)): widget.state(['disabled'])
            elif hasattr(widget,'config') and 'state' in widget.config(): widget.config(state="disabled")
        except Exception as e: print(f"Adv:No se pudo deshabilitar {widget}:{e}")
    if section_data['is_expanded'] and section_data.get('is_unlocked', False):
        section_data['is_expanded'] = False; _actualizar_estado_visual_seccion(section_data)
    try:
        current_idx = secciones_info.index(section_data)
        if current_idx + 1 < len(secciones_info):
            next_s_data = secciones_info[current_idx + 1]
            if not next_s_data.get('is_unlocked', False):
                next_s_data['is_unlocked'] = True; _actualizar_estilo_header(next_s_data)
            if next_s_data.get('paso_numero') == 2: cargar_y_mostrar_word_paso2(next_s_data)
            if not next_s_data['is_expanded']:
                next_s_data['is_expanded'] = True; _actualizar_estado_visual_seccion(next_s_data)
    except ValueError: pass

def crear_seccion_desplegable(master, titulo_sin_paso, contenido_callback, estado_inicial=False): # Modificado
    frame_externo = ttk.Frame(master); frame_externo.pack(fill=tk.X, pady=(0, 10), padx=10)
    current_section_data = {}; paso_numero = len(secciones_info) + 1
    is_initially_unlocked = (paso_numero == 1)
    header_frame = ttk.Frame(frame_externo); header_frame.pack(fill=tk.X)
    arrow_label = ttk.Label(header_frame, text="▲" if estado_inicial and is_initially_unlocked else "▼")
    arrow_label.pack(side=tk.LEFT, padx=(10, 0), pady=6)
    paso_text_label = ttk.Label(header_frame, text=f"Paso {paso_numero}")
    paso_text_label.pack(side=tk.LEFT, padx=(5,0), pady=6)
    separator_label = ttk.Label(header_frame, text="."); separator_label.pack(side=tk.LEFT, pady=6)
    title_label = ttk.Label(header_frame, text=f" {titulo_sin_paso}")
    title_label.pack(side=tk.LEFT, padx=(0,10), pady=6, fill=tk.X, expand=True)
    widgets_in_header_list = [header_frame, arrow_label, paso_text_label, separator_label, title_label]
    for widget in widgets_in_header_list:
        widget.bind("<Button-1>", lambda e, sd=current_section_data: header_click_handler(e, sd))
    frame_contenido = ttk.Frame(frame_externo, relief="groove", borderwidth=2, padding=(10, 10, 10, 0)) # padding inferior 0 para text area
    current_section_data.update({
        'paso_numero': paso_numero, # Guardar el número de paso
        'header_frame': header_frame, 'arrow_widget': arrow_label, 'widgets_in_header': widgets_in_header_list,
        'frame_contenido': frame_contenido, 'titulo_texto': f"Paso {paso_numero}. {titulo_sin_paso}",
        'is_expanded': estado_inicial and is_initially_unlocked, 'is_unlocked': is_initially_unlocked,
        'action_button_principal': None, 'ruta_imagen_seleccionada': None,
        'botones_secundarios_a_deshabilitar': [], 'widgets_contenido_a_deshabilitar': [],
        'word_title_label': None, 'word_text_widget': None # Para Paso 2
    })
    contenido_callback(frame_contenido, current_section_data)
    _actualizar_estilo_header(current_section_data)
    _actualizar_estado_visual_seccion(current_section_data)
    secciones_info.append(current_section_data)
    return current_section_data

def seleccionar_archivo_imagen(label_status_imagen, section_data_paso1, frame_contenido_paso1): # Usa PROJECT_ROOT_DIR
    extensiones_validas_display = "*.png;*.jpg;*.jpeg;*.bmp;*.tif;*.tiff"; extensiones_validas_check = ('.png', '.jpg', '.jpeg', '.bmp', '.tif', '.tiff')
    try:
        carpeta_images_inicial = os.path.join(PROJECT_ROOT_DIR, "images") # GUARDA IMAGENES EN RAIZ DEL PROYECTO
        if not os.path.exists(carpeta_images_inicial): os.makedirs(carpeta_images_inicial, exist_ok=True)
        ruta_imagen = filedialog.askopenfilename(parent=root, title="Selecciona la imagen", initialdir=carpeta_images_inicial,
            filetypes=[("Archivos de imagen", extensiones_validas_display), ("Todos los archivos", "*.*")])
        wraplength_dinamico = 580
        if not ruta_imagen:
            label_status_imagen.config(text="No se seleccionó ningún archivo.", foreground=COLOR_GRIS_TEXTO_SECUNDARIO, wraplength=wraplength_dinamico)
            section_data_paso1['ruta_imagen_seleccionada'] = None; return None
        if not ruta_imagen.lower().endswith(extensiones_validas_check):
            label_status_imagen.config(text=f"Error: Formato no válido. Use: {', '.join(ext[1:] for ext in extensiones_validas_check)}", foreground=COLOR_TEXTO_ADVERTENCIA, wraplength=wraplength_dinamico)
            section_data_paso1['ruta_imagen_seleccionada'] = None; return None
        nombre_archivo = os.path.basename(ruta_imagen)
        label_status_imagen.config(text=f"Imagen: \"{nombre_archivo}\" agregada.", foreground=COLOR_TEXTO_EXITO, wraplength=wraplength_dinamico)
        section_data_paso1['ruta_imagen_seleccionada'] = ruta_imagen; return ruta_imagen
    except Exception as e:
        messagebox.showerror("Error al Seleccionar Archivo", f"Error: {str(e)}"); label_status_imagen.config(text="Error al seleccionar.", foreground=COLOR_TEXTO_ADVERTENCIA, wraplength=wraplength_dinamico)
        section_data_paso1['ruta_imagen_seleccionada'] = None; return None

# --- Funciones de Contenido para Cada Paso ---
def contenido_paso1(frame_contenido, section_data): # Sin cambios en estructura
    label_estado_imagen = ttk.Label(frame_contenido, text="Aún no se ha seleccionado ninguna imagen.", foreground=COLOR_GRIS_TEXTO_SECUNDARIO, font=FONT_NORMAL, wraplength=580, justify="center", anchor="center")
    label_estado_imagen.pack(pady=(0,10), fill=tk.X)
    btn_agregar = tk.Button(frame_contenido, text="AGREGAR IMAGEN", relief="solid", borderwidth=1, bg=COLOR_BOTON_MINIMALISTA_BG, fg=COLOR_BOTON_MINIMALISTA_FG, activebackground=COLOR_BOTON_MINIMALISTA_ACTIVE_BG, activeforeground=COLOR_BOTON_MINIMALISTA_FG, highlightthickness=1, highlightbackground=COLOR_BOTON_MINIMALISTA_BORDER, font=FONT_NORMAL, pady=5, command=lambda: seleccionar_archivo_imagen(label_estado_imagen, section_data, frame_contenido))
    btn_agregar.pack(pady=(5,5))
    section_data['botones_secundarios_a_deshabilitar'].append(btn_agregar)
    ttk.Label(frame_contenido, text="Formatos soportados: png, jpg, jpeg, bmp, tif, tiff", foreground=COLOR_GRIS_TEXTO_SECUNDARIO, font=FONT_NORMAL, justify="center", anchor="center").pack(pady=(0,10), fill=tk.X)
    btn_procesar = tk.Button(frame_contenido, text="PROCESAR IMAGEN", relief="raised", borderwidth=1, bg=COLOR_BOTON_ACCION_PRINCIPAL_AZUL, fg=COLOR_TEXTO_BOTON_AZUL, highlightthickness=0, font=FONT_BUTTON_ACTION_MAIN)
    btn_procesar.pack(pady=(10,5), fill=tk.X, padx=20, ipady=4)
    section_data['action_button_principal'] = btn_procesar; btn_procesar.config(command=lambda: accion_principal_paso(section_data))

def contenido_paso2(frame, section_data): # Modificado para Título y Texto scrollable
    # Etiqueta para el título del Word
    label_titulo_word = ttk.Label(frame, text="Esperando documento...", font=FONT_SUB_HEADER, # Usar FONT_SUB_HEADER para el título
                                  foreground=COLOR_GRIS_TEXTO_SECUNDARIO, wraplength=600)
    label_titulo_word.pack(anchor="w", pady=(0,5), fill=tk.X)
    section_data['word_title_label'] = label_titulo_word

    # Frame para el Text y Scrollbar
    text_area_frame = ttk.Frame(frame)
    text_area_frame.pack(fill=tk.BOTH, expand=True, pady=(0,5))

    v_scrollbar = ttk.Scrollbar(text_area_frame, orient=tk.VERTICAL)
    v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    text_widget_word = tk.Text(text_area_frame, wrap=tk.WORD, yscrollcommand=v_scrollbar.set,
                               font=FONT_NORMAL, height=10, # Altura inicial en líneas
                               relief="flat", borderwidth=0,
                               padx=5, pady=5, state='disabled',
                               background=root.cget('bg')) # Fondo igual al de la ventana
    text_widget_word.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    v_scrollbar.config(command=text_widget_word.yview)
    section_data['word_text_widget'] = text_widget_word

    ttk.Label(frame, text="⚠ Si detecta errores en la transcripción, verifique el doc Word o la imagen original.",
              foreground=COLOR_TEXTO_ADVERTENCIA, wraplength=600, justify="left", font=FONT_NORMAL).pack(anchor="w", pady=(5,10), fill=tk.X)
    btn_extraer = tk.Button(frame, text="EXTRAER PARÁMETROS", relief="raised", borderwidth=1, bg=COLOR_BOTON_ACCION_PRINCIPAL_AZUL, fg=COLOR_TEXTO_BOTON_AZUL, highlightthickness=0, font=FONT_BUTTON_ACTION_MAIN)
    btn_extraer.pack(pady=5, fill=tk.X, padx=20, ipady=4)
    section_data['action_button_principal'] = btn_extraer; btn_extraer.config(command=lambda: accion_principal_paso(section_data))

def contenido_paso3(frame, section_data): # Sin cambios en estructura interna
    common_entry_width = 18
    frame_lambda = ttk.Frame(frame); frame_lambda.pack(fill=tk.X, pady=3)
    ttk.Label(frame_lambda, text="Tasa de llegadas (lambda):", font=FONT_NORMAL).pack(side="left", padx=(0,5))
    entry_lambda = ttk.Entry(frame_lambda, width=common_entry_width, font=FONT_PLACEHOLDER)
    entry_lambda.pack(side="left", padx=5); setup_placeholder(entry_lambda, PLACEHOLDER_EJ_20)
    combo_unidad_lambda = ttk.Combobox(frame_lambda, values=["pacientes/hora", "pacientes/minuto", "pacientes/día"], width=15, state="readonly", font=FONT_NORMAL)
    combo_unidad_lambda.pack(side="left", padx=5); combo_unidad_lambda.current(0)
    frame_mu = ttk.Frame(frame); frame_mu.pack(fill=tk.X, pady=3)
    ttk.Label(frame_mu, text="Tiempo de servicio (mu):", font=FONT_NORMAL).pack(side="left", padx=(0,5))
    entry_mu_valor = ttk.Entry(frame_mu, width=common_entry_width, font=FONT_PLACEHOLDER)
    entry_mu_valor.pack(side="left", padx=5); setup_placeholder(entry_mu_valor, PLACEHOLDER_INGRESE_VALOR)
    combo_unidad_mu = ttk.Combobox(frame_mu, values=["pacientes/hora", "pacientes/minuto", "minutos/paciente", "horas/paciente"], width=15, state="readonly", font=FONT_NORMAL)
    combo_unidad_mu.pack(side="left", padx=5); setup_placeholder(combo_unidad_mu, PLACEHOLDER_UNIDAD)
    frame_k = ttk.Frame(frame); frame_k.pack(fill=tk.X, pady=3)
    ttk.Label(frame_k, text="Capacidad del sistema (K):", font=FONT_NORMAL).pack(side="left", padx=(0,5))
    entry_k = ttk.Entry(frame_k, width=common_entry_width, font=FONT_PLACEHOLDER)
    entry_k.pack(side="left", padx=5); setup_placeholder(entry_k, PLACEHOLDER_INGRESE_VALOR)
    frame_s = ttk.Frame(frame); frame_s.pack(fill=tk.X, pady=3)
    ttk.Label(frame_s, text="Servidores en paralelo (s):", font=FONT_NORMAL).pack(side="left", padx=(0,5))
    entry_s = ttk.Entry(frame_s, width=common_entry_width, font=FONT_PLACEHOLDER)
    entry_s.pack(side="left", padx=5); setup_placeholder(entry_s, PLACEHOLDER_INGRESE_VALOR)
    section_data['widgets_contenido_a_deshabilitar'].extend([entry_lambda, combo_unidad_lambda, entry_mu_valor, combo_unidad_mu, entry_k, entry_s])
    btn_ingresar_modelo = tk.Button(frame, text="INGRESAR DATOS AL MODELO", relief="raised", borderwidth=1, bg=COLOR_BOTON_ACCION_PRINCIPAL_AZUL, fg=COLOR_TEXTO_BOTON_AZUL, highlightthickness=0, font=FONT_BUTTON_ACTION_MAIN)
    btn_ingresar_modelo.pack(pady=(15,5), fill=tk.X, padx=20, ipady=4)
    section_data['action_button_principal'] = btn_ingresar_modelo; btn_ingresar_modelo.config(command=lambda: accion_principal_paso(section_data))

def contenido_paso4(frame, section_data): # Sin cambios en estructura interna
    ttk.Label(frame, text="Los datos corresponden a un modelo M/M/1/K (rho != 1)", foreground=COLOR_TEXTO_EXITO, font=FONT_SECTION_HEADER, justify="center", anchor="center").pack(pady=(0,10), fill=tk.X)
    ttk.Label(frame, text="Seleccione que desea calcular:", justify="center", anchor="center", font=FONT_NORMAL).pack(pady=(0,8), fill=tk.X)
    frame_botones_calculo = ttk.Frame(frame); frame_botones_calculo.pack(pady=5)
    btn_medidas = tk.Button(frame_botones_calculo, text="Medidas de desempeño", relief="raised", borderwidth=1, pady=3, padx=10, font=FONT_NORMAL, bg=COLOR_BOTON_MINIMALISTA_BG, fg=COLOR_BOTON_MINIMALISTA_FG)
    btn_medidas.pack(side="left", padx=5)
    btn_probabilidad = tk.Button(frame_botones_calculo, text="Probabilidad de \"n\" clientes", relief="raised", borderwidth=1, pady=3, padx=10, font=FONT_NORMAL, bg=COLOR_BOTON_MINIMALISTA_BG, fg=COLOR_BOTON_MINIMALISTA_FG)
    btn_probabilidad.pack(side="left", padx=5)

# --- Crear Ventana Principal y Estilos ---
root = tk.Tk()
root.title("TITULO DEL PROYECTO"); root.minsize(720, 780); root.configure(bg="#f0f0f0")
style = ttk.Style()
style.configure("HeaderFrame.TFrame", background=COLOR_HEADER_ACTIVE_BG, relief="raised", borderwidth=1)
style.configure("HeaderLabel.TLabel", background=COLOR_HEADER_ACTIVE_BG, foreground=COLOR_HEADER_ACTIVE_FG)
style.configure("LockedHeaderFrame.TFrame", background=COLOR_HEADER_LOCKED_BG, relief="flat", borderwidth=1)
style.configure("LockedHeaderLabel.TLabel", background=COLOR_HEADER_LOCKED_BG, foreground=COLOR_HEADER_LOCKED_FG)
style.configure('.', font=FONT_NORMAL); style.configure('TButton', font=FONT_NORMAL)
style.configure('TLabel', font=FONT_NORMAL); style.configure('TEntry', font=FONT_PLACEHOLDER)
style.configure('TCombobox', font=FONT_NORMAL)
root.option_add("*TCombobox*Listbox*Font", FONT_NORMAL)
titulo_label = ttk.Label(root, text="TITULO DEL PROYECTO", font=FONT_MAIN_TITLE, background="#f0f0f0")
titulo_label.pack(pady=(15, 25))

# --- Crear Pasos Desplegables ---
crear_seccion_desplegable(root, "Adjuntar imagen.", contenido_paso1, estado_inicial=True)
crear_seccion_desplegable(root, "Vista previa del documento WORD #1 generado.", contenido_paso2)
crear_seccion_desplegable(root, "Datos extraídos.", contenido_paso3)
crear_seccion_desplegable(root, "Modelo de teoría de colas.", contenido_paso4)

root.mainloop()